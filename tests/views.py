from django.shortcuts import render, redirect, get_object_or_404
from django.http import Http404, HttpResponse
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from docx import Document
from docx.shared import Inches
from tests.forms import (
    CreateTestForm, 
    CreateResultForm,
    VariantFormSet,
    QuestionForm,
    SetQuestionsForm,
    SetTestForm,
)
from tests.decorators import psych_required
from tests.models import (
    Category, 
    Test,
    Status, 
    AssigningTest, 
    Question,
    TestResult
)
from collections import Counter
from django.db.models import Avg, Q, Max, F
from tests.utils import custom_get_object
from users.models import User, Personality, Tendency, Activity, UserGroup


def index(request):
    if request.user.is_authenticated and request.user.role.name == 'студент':
        assigning_tests = AssigningTest.objects.filter(user=request.user).select_related('test')
        tests = [assigning_test.test for assigning_test in assigning_tests]

        test_ids = [test.id for test in tests]

        completed_test_ids = set(
            TestResult.objects.filter(user=request.user, test_id__in=test_ids)
            .values_list('test_id', flat=True)
        )

        for test in tests:
            test.is_completed = test.id in completed_test_ids
    else:
        tests = Test.objects.all()
    
    personality_types = Personality.objects.all()

    return render(request, "tests/index.html", {
        'personality_types': personality_types,
        'tests': tests,
    })
    
@psych_required()
def test_results_chart(request):
    tests = Test.objects.all()

    labels = []
    avg_scores = []

    for test in tests:
        labels.append(test.name)
        avg_score = TestResult.objects.filter(test=test).aggregate(avg=Avg('scores'))['avg']
        avg_scores.append(round(avg_score, 2) if avg_score else 0)

    context = {
        'labels': labels,
        'scores': avg_scores,
    }

    return render(request, 'tests/test_results_chart.html', context)

@psych_required()
def students_list(request):
    students = User.objects_students.all()
    
    trial_test = Test.objects.filter(trial=True).first()

    search_query = request.GET.get('search', '')
    group_id = request.GET.get('group')
    personality_id = request.GET.get('personality')
    sort = request.GET.get('sort', 'last_name')  # по умолчанию сортируем по фамилии
    direction = request.GET.get('direction', 'asc')

    if search_query:
        students = students.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(patronymic__icontains=search_query)
        )

    if group_id:
        students = students.filter(group__id=group_id)
    if personality_id:
        students = students.filter(personality_type__id=personality_id)

    if direction == 'desc':
        sort = f'-{sort}'
    students = students.order_by(sort)

    personality_types = students.values_list('personality_type__name', flat=True)
    personality_counts = Counter(personality_types)
    labels = list(personality_counts.keys())
    data = list(personality_counts.values())

    groups = UserGroup.objects.all()
    personalities = Personality.objects.all()

    return render(request, 'tests/students_list.html', {
        'students': students,
        'labels': labels,
        'data': data,
        'groups': groups,
        'personalities': personalities,
        'current_sort': request.GET.get('sort', 'last_name'),
        'current_direction': direction,
        'trial_test': trial_test,
    })

@psych_required()
def personal_type(request, personality_id):
    personality = get_object_or_404(Personality, id=personality_id)
    students = User.objects.filter(personality_type=personality)
    tests = Test.objects.all()

    statuses = {}

    for student in students:
        statuses[student.pk] = {}
        
        for test in tests:
            test_results = TestResult.objects.filter(test=test, user=student)
            assign_test = AssigningTest.objects.filter(test=test, user=student).first()

            if test_results.exists():
                statuses[student.pk][test.pk] = 'завершен'
            elif assign_test and timezone.now() >= assign_test.due_datetime:
                statuses[student.pk][test.pk] = 'не завершен'
            elif assign_test:
                statuses[student.pk][test.pk] = 'в процессе'
            else:
                statuses[student.pk][test.pk] = 'не назначен'
                
    return render(request, 'tests/personal_type.html', {
        'personality': personality,
        'students': students,
        'tests': tests,
        'statuses': statuses,
    })

@login_required(login_url='/users/login/')
def show_test(request, test_id):
    test = custom_get_object(request, test_id)
    count_questions = Question.objects.filter(test=test).count()
    
    return render(request, "tests/test.html", {
        'count_questions': count_questions,
        'test': test
    })

@login_required(login_url='/users/login/')
def test_result(request, test_id):
    test_result = get_object_or_404(TestResult, test=test_id, user=request.user)
    return render(request, 'tests/test_result.html', {'test_result': test_result.result})

@login_required(login_url='/users/login/')
def questions_test(request, test_id):
    test = get_object_or_404(Test, id=test_id)
    assign_test = AssigningTest.objects.filter(user=request.user, test=test).first()
    if (test.trial and TestResult.objects.filter(test=test, user=request.user).exists()) or \
    (not assign_test):
        raise Http404()
    
    questions = Question.objects.filter(test=test).prefetch_related('variant_set')

    if request.method == 'POST':
        form = SetQuestionsForm(
            request.POST, 
            questions=questions, 
            test=test, 
            user=request.user,
            assign_test=assign_test
        )
        if form.is_valid():
            form.save()
            return redirect('tests:test_result', test_id=test.pk)
    else:
        form = SetQuestionsForm(
            questions=questions, 
            test=test, 
            user=request.user,
            assign_test=assign_test
        )

    return render(request, 'tests/questions_test.html', {
        'test': test,
        'form': form,
        'questions': questions,
    })

@psych_required()
def questions_list(request, test_id):
    test = get_object_or_404(Test, id=test_id)
    questions = Question.objects.filter(test=test)
    return render(request, "tests/questions_list.html", {'questions': questions})

@psych_required()
def add_question(request, test_id):
    test = get_object_or_404(Test, id=test_id)

    if request.method == 'POST':
        form = QuestionForm(request.POST)
        formset = VariantFormSet(request.POST)

        if form.is_valid() and formset.is_valid():
            question = form.save(commit=False)
            question.test = test
            last_order = Question.objects.filter(test=test).aggregate(Max('order'))['order__max'] or 1
            question.order = last_order + 1
            question.save()
            formset.instance = question
            formset.save()
            
            return redirect('home')

    else:
        form = QuestionForm()
        formset = VariantFormSet()

    return render(request, 'tests/add_question.html', {
        'form': form,
        'formset': formset,
        'test': test
    })
    
@psych_required()
def update_question(request, question_id):
    question = get_object_or_404(Question, id=question_id)
    test = question.test
    current_order = question.order

    if request.method == 'POST':
        form = QuestionForm(request.POST, instance=question)
        formset = VariantFormSet(request.POST, instance=question)

        if form.is_valid() and formset.is_valid():
            updated_question = form.save(commit=False)
            new_order = updated_question.order

            questions_qs = Question.objects.filter(test=test).exclude(id=question.id).order_by('order')

            max_order = questions_qs.aggregate(Max('order'))['order__max'] or 0

            if new_order is None or new_order < 1:
                new_order = 1
            elif new_order > max_order + 1:
                new_order = max_order + 1

            if current_order and new_order != current_order:
                if new_order < current_order:
                    Question.objects.filter(
                        test=test,
                        order__gte=new_order,
                        order__lt=current_order
                    ).update(order=F('order') + 1)
                elif new_order > current_order:
                    Question.objects.filter(
                        test=test,
                        order__gt=current_order,
                        order__lte=new_order
                    ).update(order=F('order') - 1)

            updated_question.order = new_order
            updated_question.save()
            formset.save()
            return redirect('home')
    else:
        form = QuestionForm(instance=question)
        formset = VariantFormSet(instance=question)

    return render(request, 'tests/update_question.html', {
        'form': form,
        'formset': formset,
        'test': test,
        'question': question,
    })
    
@psych_required()
def delete_question(request, question_id):
    question = get_object_or_404(Question, id=question_id)
    test = question.test
    deleted_order = question.order

    question.delete()

    if deleted_order is not None:
        Question.objects.filter(
            test=test,
            order__gt=deleted_order,
            order__isnull=False
        ).update(order=F('order') - 1)

    return redirect('tests:questions_list', test_id=test.pk)

@psych_required()
def add_result(request, test_id):
    test = get_object_or_404(Test, id=test_id)
    activities = Activity.objects.all()
    tendencies = Tendency.objects.all()
    
    if request.method == 'POST':
        form = CreateResultForm(request.POST)
        if form.is_valid():
            result = form.save(commit=False)
            result.test = test            
            result.save()
            form.save_m2m()
            
            return redirect('home')
    else:
        form = CreateResultForm()
    
    return render(request, "tests/add_result.html", {
        'form': form,
        'activities': activities,
        'tendencies': tendencies,
        'test_pk': test.pk,
    })

@psych_required()
def assign_test(request, test_id):
    test = get_object_or_404(Test, id=test_id)
    
    if request.method == 'POST':
        form = SetTestForm(request.POST)
        if form.is_valid():
            result = form.save(commit=False)
            result.author = request.user
            result.test = test
            
            status, _ = Status.objects.get_or_create(name="в процессе")
            result.status = status
            
            result.save()

            return redirect('home')
    
    if test.trial:
        students = User.objects_students.filter(personality_type=None)
    else:
        students = User.objects_students.all()
    
    return render(request, 'tests/assign_test.html', {
        'students': students,
        'test_pk': test.pk,
    })

@psych_required()
def create(request):
    categories = Category.objects.all()
    if request.method == 'POST':
        form = CreateTestForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('home')
    else:
        form = CreateTestForm()
    return render(request, 'tests/create.html', {'form': form, 'categories': categories})

import io
from datetime import datetime

def generate_student_report(request, personality_id):
    """Generates a Word report of student data for a personality type."""

    personality = Personality.objects.get(pk=personality_id)
    students = User.objects.filter(role='2')  # Retrieve only students
    tests = Test.objects.all()
    statuses = get_statuses_data(students, tests)

    # Create a new Word document
    document = Document()

    # Add report title and metadata
    document.add_paragraph(f"Дата создания: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph(f"Психолог: {request.user.first_name} {request.user.last_name}")

    # Create the table
    table = document.add_table(rows=1, cols=len(tests) + 1)
    header_cells = table.rows[0].cells
    header_cells[0].text = "ФИО"
    for i, test in enumerate(tests):
        header_cells[i + 1].text = test.name

    # Populate the table with student data
    for student in students:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{student.last_name} {student.first_name} {student.patronymic}"

        for i, test in enumerate(tests):
            status = statuses.get(student.pk, {}).get(test.pk, "N/A")
            row_cells[i + 1].text = status

    # Adjust table column widths (optional)
    for i in range(len(tests) + 1):
        for cell in table.columns[i].cells:
            cell.width = Inches(1.0)

    # Save the document to a buffer
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # Create the HTTP response
    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=student_report_{personality.name}.docx'
    return response

def get_statuses_data(students, tests):
    statuses = {}
    for student in students:
        student_statuses = {}
        for test in tests:
            assigning_test = AssigningTest.objects.filter(user=student, test=test).first()
            if assigning_test:
                student_statuses[test.pk] = assigning_test.status.name
            else:
                student_statuses[test.pk] = "N/A"
        statuses[student.pk] = student_statuses
    return statuses



def students_lists(request, personality_id):
    personality = get_object_or_404(Personality, pk=personality_id)
    students = User.objects.filter(role='student')  # Retrieve only students
    tests = Test.objects.all()
    statuses = get_statuses_data()
    return render(request, 'tests/student_list.html', {
        'personality': personality,
        'students': students,
        'tests': tests,
        'statuses': statuses,
    })

    
    